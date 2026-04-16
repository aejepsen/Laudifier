import tempfile
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors


class ExportService:
    async def exportar(self, laudo: dict, formato: str) -> str:
        texto = laudo.get("laudo_editado") or laudo.get("laudo", "")
        if formato == "pdf":
            return self._to_pdf(texto, laudo)
        elif formato == "docx":
            return self._to_docx(texto, laudo)
        else:
            return self._to_txt(texto, laudo)

    def _to_pdf(self, texto: str, laudo: dict) -> str:
        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        doc = SimpleDocTemplate(tmp.name, pagesize=A4,
                                leftMargin=2.5*cm, rightMargin=2.5*cm,
                                topMargin=2.5*cm, bottomMargin=2.5*cm)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("titulo", parent=styles["Heading1"],
                                     fontSize=14, textColor=colors.HexColor("#1a56db"))
        body_style  = ParagraphStyle("corpo", parent=styles["Normal"],
                                     fontSize=11, leading=16)
        elements = [
            Paragraph(f"LAUDO MÉDICO — {laudo.get('especialidade','').upper()}", title_style),
            Spacer(1, 0.5*cm),
            Paragraph(f"Data: {laudo.get('created_at','')[:10]}", styles["Normal"]),
            Spacer(1, 0.3*cm),
        ]
        for linha in texto.split("\n"):
            if linha.strip():
                if linha.isupper() and len(linha) < 60:
                    elements.append(Paragraph(linha, styles["Heading2"]))
                else:
                    elements.append(Paragraph(linha.replace("**", ""), body_style))
                elements.append(Spacer(1, 0.1*cm))
        doc.build(elements)
        return tmp.name

    def _to_docx(self, texto: str, laudo: dict) -> str:
        from docx import Document as DocxDoc
        doc  = DocxDoc()
        doc.add_heading(f"LAUDO — {laudo.get('especialidade','').upper()}", 0)
        doc.add_paragraph(f"Data: {laudo.get('created_at','')[:10]}")
        doc.add_paragraph("")
        for linha in texto.split("\n"):
            if linha.strip():
                p = doc.add_paragraph(linha.replace("**", ""))
                if linha.isupper() and len(linha) < 60:
                    p.style = "Heading 2"
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        return tmp.name

    def _to_txt(self, texto: str, laudo: dict) -> str:
        tmp = tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w", encoding="utf-8")
        tmp.write(f"LAUDO MÉDICO — {laudo.get('especialidade','').upper()}\n")
        tmp.write(f"Data: {laudo.get('created_at','')[:10]}\n\n")
        tmp.write(texto)
        tmp.close()
        return tmp.name
