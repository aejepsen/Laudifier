# backend/services/laudo_service.py
"""Persistência de laudos no Supabase."""
import os
from datetime import datetime, timezone
from supabase import create_client

SB_URL = os.getenv("SUPABASE_URL")
SB_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")

# Schema Supabase (execute no SQL Editor):
SUPABASE_SCHEMA = """
CREATE TABLE IF NOT EXISTS laudos (
    id             UUID PRIMARY KEY DEFAULT gen_random_uuid(),
    user_id        UUID REFERENCES auth.users(id) ON DELETE CASCADE,
    especialidade  TEXT NOT NULL,
    tipo_laudo     TEXT,
    solicitacao    TEXT,
    laudo          TEXT NOT NULL,
    laudo_editado  TEXT,
    tipo_geracao   TEXT,     -- 'rag' | 'fallback'
    laudos_ref     JSONB DEFAULT '[]',
    aprovado       BOOLEAN,
    correcoes      TEXT,
    created_at     TIMESTAMPTZ DEFAULT NOW(),
    updated_at     TIMESTAMPTZ DEFAULT NOW()
);
ALTER TABLE laudos ENABLE ROW LEVEL SECURITY;
CREATE POLICY "users_own_laudos" ON laudos FOR ALL USING (auth.uid() = user_id);
CREATE INDEX idx_laudos_especialidade ON laudos(especialidade);
CREATE INDEX idx_laudos_user_created ON laudos(user_id, created_at DESC);

CREATE TABLE IF NOT EXISTS user_profiles (
    id             UUID PRIMARY KEY DEFAULT gen_random_uuid(),
    user_id        UUID REFERENCES auth.users(id) ON DELETE CASCADE UNIQUE,
    display_name   TEXT,
    crm            TEXT,
    especialidade  TEXT,
    role           TEXT DEFAULT 'medico',
    created_at     TIMESTAMPTZ DEFAULT NOW()
);
ALTER TABLE user_profiles ENABLE ROW LEVEL SECURITY;
CREATE POLICY "users_own_profile" ON user_profiles FOR ALL USING (auth.uid() = user_id);
"""

class LaudoService:
    def __init__(self, user_id: str):
        self.user_id = user_id
        self.sb = create_client(SB_URL, SB_KEY)

    async def salvar(self, laudo_id, especialidade, solicitacao, laudo, tipo_geracao, laudos_ref):
        self.sb.table("laudos").upsert({
            "id": laudo_id, "user_id": self.user_id,
            "especialidade": especialidade, "solicitacao": solicitacao,
            "laudo": laudo, "tipo_geracao": tipo_geracao,
            "laudos_ref": laudos_ref,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }).execute()

    async def listar(self, page=0, size=20, especialidade=None):
        q = self.sb.table("laudos").select(
            "id, especialidade, tipo_laudo, solicitacao, tipo_geracao, aprovado, created_at, updated_at"
        ).eq("user_id", self.user_id).order("created_at", desc=True).range(page*size, (page+1)*size-1)
        if especialidade:
            q = q.eq("especialidade", especialidade)
        return (q.execute()).data or []

    async def get(self, laudo_id: str) -> dict | None:
        r = self.sb.table("laudos").select("*").eq("id", laudo_id).eq("user_id", self.user_id).single().execute()
        return r.data

    async def atualizar(self, laudo_id: str, laudo_editado: str):
        self.sb.table("laudos").update({
            "laudo_editado": laudo_editado,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }).eq("id", laudo_id).eq("user_id", self.user_id).execute()

    async def registrar_feedback(self, laudo_id: str, aprovado: bool, correcoes: str | None):
        self.sb.table("laudos").update({
            "aprovado": aprovado, "correcoes": correcoes,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }).eq("id", laudo_id).eq("user_id", self.user_id).execute()

    async def get_stats(self) -> dict:
        # Carrega apenas as colunas necessárias com LIMIT explícito.
        # Supabase não suporta GROUP BY nativo na client lib,
        # mas limitamos a 1000 registros para evitar full table scan irrestrito.
        dados = (
            self.sb.table("laudos")
            .select("especialidade, tipo_geracao, aprovado")
            .eq("user_id", self.user_id)
            .limit(1000)
            .execute()
            .data or []
        )
        total = len(dados)
        rag   = sum(1 for l in dados if l.get("tipo_geracao") == "rag")
        aprov = sum(1 for l in dados if l.get("aprovado") is True)
        by_esp: dict[str, int] = {}
        for l in dados:
            esp = l.get("especialidade") or "outros"
            by_esp[esp] = by_esp.get(esp, 0) + 1
        return {
            "total_laudos":      total,
            "por_rag":           rag,
            "por_fallback":      total - rag,
            "aprovados":         aprov,
            "taxa_aprovacao":    round(aprov / total, 2) if total else 0,
            "por_especialidade": by_esp,
        }

    async def deletar(self, laudo_id: str) -> None:
        """Remove laudo do usuário (LGPD art. 18, VI — direito de exclusão)."""
        self.sb.table("laudos") \
            .delete() \
            .eq("id", laudo_id) \
            .eq("user_id", self.user_id) \
            .execute()


# ─────────────────────────────────────────────────────────────────────────────
# backend/services/export_service.py
# ─────────────────────────────────────────────────────────────────────────────

import tempfile
from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors

class ExportService:
    async def exportar(self, laudo: dict, formato: str) -> str:
        texto = laudo.get("laudo_editado") or laudo.get("laudo", "")
        if formato == "pdf":
            return self._to_pdf(texto, laudo)
        elif formato == "docx":
            return self._to_docx(texto, laudo)
        else:
            return self._to_txt(texto, laudo)

    def _to_pdf(self, texto: str, laudo: dict) -> str:
        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        doc = SimpleDocTemplate(tmp.name, pagesize=A4,
                                leftMargin=2.5*cm, rightMargin=2.5*cm,
                                topMargin=2.5*cm, bottomMargin=2.5*cm)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("titulo", parent=styles["Heading1"],
                                     fontSize=14, textColor=colors.HexColor("#1a56db"))
        body_style  = ParagraphStyle("corpo", parent=styles["Normal"],
                                     fontSize=11, leading=16)
        elements = [
            Paragraph(f"LAUDO MÉDICO — {laudo.get('especialidade','').upper()}", title_style),
            Spacer(1, 0.5*cm),
            Paragraph(f"Data: {laudo.get('created_at','')[:10]}", styles["Normal"]),
            Spacer(1, 0.3*cm),
        ]
        for linha in texto.split("\n"):
            if linha.strip():
                if linha.isupper() and len(linha) < 60:
                    elements.append(Paragraph(linha, styles["Heading2"]))
                else:
                    elements.append(Paragraph(linha.replace("**", ""), body_style))
                elements.append(Spacer(1, 0.1*cm))
        doc.build(elements)
        return tmp.name

    def _to_docx(self, texto: str, laudo: dict) -> str:
        from docx import Document as DocxDoc
        doc  = DocxDoc()
        doc.add_heading(f"LAUDO — {laudo.get('especialidade','').upper()}", 0)
        doc.add_paragraph(f"Data: {laudo.get('created_at','')[:10]}")
        doc.add_paragraph("")
        for linha in texto.split("\n"):
            if linha.strip():
                p = doc.add_paragraph(linha.replace("**", ""))
                if linha.isupper() and len(linha) < 60:
                    p.style = "Heading 2"
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        return tmp.name

    def _to_txt(self, texto: str, laudo: dict) -> str:
        tmp = tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w", encoding="utf-8")
        tmp.write(f"LAUDO MÉDICO — {laudo.get('especialidade','').upper()}\n")
        tmp.write(f"Data: {laudo.get('created_at','')[:10]}\n\n")
        tmp.write(texto)
        tmp.close()
        return tmp.name


# ─────────────────────────────────────────────────────────────────────────────
# backend/services/prompt_service.py
# ─────────────────────────────────────────────────────────────────────────────

import functools
from pathlib import Path

PROMPT_DIR = Path(__file__).parent.parent / "prompts"

@functools.lru_cache(maxsize=1)
def load_system_prompt() -> str:
    return (PROMPT_DIR / "system_prompt.txt").read_text(encoding="utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# backend/services/storage_service.py  (reutilizado do AI Professor v2)
# ─────────────────────────────────────────────────────────────────────────────

import uuid
import boto3
from botocore.config import Config
from pathlib import Path as _Path

_USE_LOCAL  = os.getenv("USE_LOCAL_STORAGE", "true").lower() == "true"
_LOCAL_PATH = _Path(os.getenv("LOCAL_STORAGE_PATH", "./storage"))

class StorageService:
    def __init__(self):
        if _USE_LOCAL:
            (_LOCAL_PATH / "laudos").mkdir(parents=True, exist_ok=True)
            (_LOCAL_PATH / "exports").mkdir(exist_ok=True)
            self._s3 = None
        else:
            self._s3 = boto3.client(
                "s3",
                endpoint_url=os.getenv("S3_ENDPOINT_URL"),
                aws_access_key_id=os.getenv("S3_ACCESS_KEY_ID"),
                aws_secret_access_key=os.getenv("S3_SECRET_ACCESS_KEY"),
                region_name=os.getenv("S3_REGION", "auto"),
                config=Config(signature_version="s3v4"),
            )

    @staticmethod
    def _sanitize_filename(filename: str) -> str:
        """Remove path separators e caracteres perigosos — previne path traversal."""
        import re
        safe = re.sub(r"[^\w\-_\.]", "_", _Path(filename).name)
        return safe[:100] or "upload"

    def upload_document(self, content: bytes, filename: str, user_id: str) -> str:
        safe_name = self._sanitize_filename(filename)
        key = f"{user_id}/{uuid.uuid4().hex}/{safe_name}"
        if _USE_LOCAL:
            path = _LOCAL_PATH / "laudos" / key
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_bytes(content)
            return str(path)
        bucket = os.getenv("S3_BUCKET_LAUDOS", "laudifier-laudos")
        self._s3.put_object(Bucket=bucket, Key=key, Body=content)
        return f"{os.getenv('S3_ENDPOINT_URL')}/{bucket}/{key}"
